"""Render the Technical Design Document (TDD) and embed the user-story cards.

Same visual language as docs/Research_Report — narrative prose, styled tables,
auto-numbered figures, a Word TOC field. Figures come from
additional-docs/screenshots/ (the styled C4 / sequence / dataset renders and the
six user-story epic captures).

    .venv/Scripts/python.exe -m scripts.build_tdd_docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SHOTS = ROOT / "additional-docs" / "screenshots"
OUT = ROOT / "additional-docs" / "Technical_Design_Document.docx"

INK = RGBColor(0x1F, 0x2A, 0x44)
MUTED = RGBColor(0x55, 0x65, 0x70)
_FIG = 0


# --------------------------------------------------------------------------- helpers
def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.25
    for h, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        s = doc.styles[h]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = INK
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(6)


def _shade(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(12)


def _figure(doc, filename: str, caption: str, width_cm: float = 15.0) -> None:
    global _FIG
    _FIG += 1
    path = SHOTS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    _caption(doc, f"Figure {_FIG} — {caption}")


def _p(doc, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if "**" in it:
            for i, seg in enumerate(it.split("**")):
                run = p.add_run(seg)
                run.bold = i % 2 == 1
        else:
            p.add_run(it)


def _table(doc, header, rows, widths=None, fill="1F2A44", fs=9.5) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = "Light List Accent 1"
    table.autofit = widths is None
    for ri, r in enumerate([header] + rows):
        for ci, val in enumerate(r):
            cell = table.cell(ri, ci)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(fs)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths is not None:
                cell.width = Cm(widths[ci])
            if ri == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade(cell, fill)
            elif ri % 2 == 0:
                _shade(cell, "EEF1F8")
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        th = OxmlElement("w:trHeight")
        th.set(qn("w:val"), "260")
        trPr.append(th)
    doc.add_paragraph()


def _h1(doc, t): doc.add_heading(t, level=1)
def _h2(doc, t): doc.add_heading(t, level=2)
def _page_break(doc): doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _toc(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    ph = OxmlElement("w:t"); ph.text = "Right-click and choose “Update Field” to build the table of contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (b, instr, sep, ph, end):
        run._r.append(el)


def _meta_table(doc, rows) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for (k, v), row in zip(rows, table.rows):
        row.cells[0].text = k
        row.cells[1].text = v
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


# --------------------------------------------------------------------------- content
def build() -> None:
    doc = Document()
    _set_default_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Network Traffic Anomaly Detector")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = INK
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Technical Design Document")
    r.italic = True; r.font.size = Pt(13); r.font.color.rgb = MUTED
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("User stories, architecture, data, API and runtime design")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = MUTED
    doc.add_paragraph()

    _meta_table(doc, [
        ("Student", "Angel Rusev"),
        ("Programme", "BSc ICT & Software Engineering"),
        ("Minor", "Cybersecurity — Attack & Defend"),
        ("Institution", "Fontys University of Applied Sciences"),
        ("Semester", "Spring 2026"),
        ("Document", "Technical Design Document v1.0 — companion to the Research Report"),
        ("Repository", "https://github.com/Guts1313/AnomalyDetector"),
    ])
    doc.add_paragraph()
    _h1(doc, "Table of contents")
    _toc(doc)
    _page_break(doc)

    # 1. Introduction
    _h1(doc, "1. Introduction and scope")
    _p(doc,
       "This Technical Design Document (TDD) is the engineering companion to the Research Report. "
       "Where the research document argues why the detector is built the way it is, this document "
       "specifies what is built: the user stories the system satisfies, the architecture that "
       "realises them, the data it handles, the API it exposes, and the runtime behaviour that "
       "ties the pieces together. It is written to be read by an engineer who has to extend, "
       "operate, or review the system, and every design statement is traceable to a user story and "
       "to code in the accompanying repository.")
    _p(doc,
       "The scope is the detector itself — a stateless FastAPI inference service, a serialised "
       "model bundle, a SQLite audit store, two analyst front-ends (React and Streamlit), and the "
       "containerised attack/defend laboratory that drives real offensive traffic through the same "
       "inference path. Model research, dataset generation and the per-class evaluation are "
       "covered in the Research Report and only summarised here where the design depends on them.")

    # 2. Personas
    _h1(doc, "2. Personas and stakeholders")
    _p(doc,
       "The design serves four personas. The SOC analyst is the primary user of the running "
       "system; the ML engineer maintains the model; the NOC operator keeps the service healthy; "
       "and the security student or researcher drives the attack/defend laboratory. Table 1 maps "
       "each persona to the goals the design must satisfy.")
    _table(doc,
        ["Persona", "Goal", "Touch-points"],
        [
            ["SOC analyst", "Triage, investigate and explain alerts in real time.", "React/Streamlit dashboards, /alerts, /predictions"],
            ["ML engineer", "Retrain, evaluate and hot-reload the model safely.", "scripts/train, /admin/reload, /metrics"],
            ["NOC operator", "Keep the service healthy and detect drift.", "/health, /metrics, Docker healthchecks"],
            ["Security student / researcher", "Exercise the detector with real attack traffic.", "Examples tab, lab attacker/defender containers"],
        ],
        widths=[3.6, 6.0, 5.4])
    _caption(doc, "Table 1 — Personas and the goals the design must satisfy.")

    # 3. User stories
    _h1(doc, "3. User stories")
    _p(doc,
       "The backlog is organised into six epics, prioritised with MoSCoW and sized in story "
       "points. The cards below are the canonical statements; each carries a role, a "
       "“so that” benefit, Given/When/Then acceptance criteria, and links to the "
       "functional requirements (FR) and sub-research questions (SRQ) it traces to. The same FR/SRQ "
       "tags reappear in the traceability matrix in §13.")
    _figure(doc, "us_epic_a.png", "Epic A — Real-time detection (US-01, US-02).")
    _figure(doc, "us_epic_b.png", "Epic B — Investigation and audit (US-03, US-04).")
    _figure(doc, "us_epic_c.png", "Epic C — Tuning and explainability (US-05, US-06).")
    _figure(doc, "us_epic_d.png", "Epic D — Model lifecycle (US-07, US-08).")
    _figure(doc, "us_epic_e.png", "Epic E — Attack / defend lab (US-09, US-10).")
    _figure(doc, "us_epic_f.png", "Epic F — Operations, security and compliance (US-11, US-12).")
    _p(doc,
       "Table 2 summarises the backlog so it can be read at a glance and sorted by priority. The "
       "Must-have stories form the walking skeleton — real-time scoring, severity triage, the "
       "audit trail, hot-reload, monitoring and the privacy guarantee — and were delivered first; "
       "the Could-have lab stories were the final, highest-value-per-risk increment.")
    _table(doc,
        ["ID", "Story (abridged)", "Role", "Priority", "Pts"],
        [
            ["US-01", "Score incoming flows in real time", "SOC analyst", "Must", "8"],
            ["US-02", "Colour-code alerts by severity", "SOC analyst", "Must", "3"],
            ["US-03", "Queryable audit log of every classification", "SOC analyst", "Must", "5"],
            ["US-04", "Pivot on source/destination IPs", "SOC analyst", "Should", "3"],
            ["US-05", "Decision-threshold slider, no retrain", "SOC analyst", "Should", "3"],
            ["US-06", "Manual scoring with class probabilities", "SOC analyst", "Should", "5"],
            ["US-07", "Retrain and hot-reload without downtime", "ML engineer", "Must", "5"],
            ["US-08", "Per-class evaluation report", "ML engineer", "Should", "5"],
            ["US-09", "Launch a real attack from the UI", "Security student", "Could", "8"],
            ["US-10", "Auto-block a confirmed attacker IP", "Defender", "Could", "8"],
            ["US-11", "Health and metrics endpoints", "NOC operator", "Must", "3"],
            ["US-12", "No payloads stored (GDPR)", "Security officer", "Must", "3"],
        ],
        widths=[1.6, 6.8, 3.0, 2.0, 1.2])
    _caption(doc, "Table 2 — Product backlog summary (MoSCoW + story points).")

    _page_break(doc)

    # 4. Architecture
    _h1(doc, "4. Architecture overview")
    _p(doc,
       "The system is a small set of cooperating containers around one inference service. Figure 7 "
       "is the C4 level-1 context: the detector sits between offline flow-feature sources and the "
       "human roles it serves, and can forward a severity feed to an upstream SIEM. Figure 8 opens "
       "the box one level further into containers.")
    _figure(doc, "c4_context.png",
            "C4 level-1 context — the detector between flow sources, the analyst, the ML engineer "
            "and an optional SIEM.")
    _figure(doc, "c4_container.png",
            "C4 level-2 containers — the stateless FastAPI hub, its model bundle and SQLite store, "
            "the React and Streamlit clients, and the attacker/defender lab.")
    _p(doc, "Three architectural rules are load-bearing and are referenced throughout the design:")
    _bullets(doc, [
        "**Stateless inference** — /predict holds no per-request state; the audit write happens after the verdict, off the critical path (US-01).",
        "**One serialised bundle** — the fitted preprocessing pipeline and the model travel together in best.joblib, so train-time and serve-time transforms cannot diverge (US-07).",
        "**Batch-first** — /predict always accepts a list of flows so the vectorised scikit-learn path does the work (US-01).",
    ])

    # 5. Components
    _h1(doc, "5. Component design")
    _p(doc,
       "Table 3 lists the runtime components, their responsibility, and the primary user stories "
       "each one realises. The boundaries are deliberately sharp: the API never renders UI, the "
       "dashboards never touch the model, and the lab containers reach the model only through the "
       "same public /predict contract a real client would use.")
    _table(doc,
        ["Component", "Responsibility", "Realises"],
        [
            ["FastAPI service (api/main.py)", "HTTP surface, request validation, verdict + severity, audit write.", "US-01, US-02, US-11"],
            ["ModelRegistry (models/registry)", "Load best.joblib; hot-swap on /admin/reload.", "US-07"],
            ["FeaturePipeline (features/)", "Median impute → RobustScaler → one-hot; drop IPs/ports.", "US-01, US-12"],
            ["AlertStore (api/store.py)", "SQLite persistence of every classification.", "US-03, US-04"],
            ["React SPA / Streamlit", "Severity-coloured dashboards, manual scoring, examples.", "US-02, US-05, US-06"],
            ["Lab attacker (lab/attacker)", "Run nmap/hping3/hydra/sqlmap behind /attack.", "US-09"],
            ["Lab defender (lab/defender)", "Capture, extract, classify, iptables auto-block.", "US-09, US-10"],
        ],
        widths=[4.6, 6.8, 3.4])
    _caption(doc, "Table 3 — Components, responsibilities and the user stories they realise.")

    # 6. Data design
    _h1(doc, "6. Data design")
    _p(doc,
       "The model consumes a fixed twenty-feature flow vector and nothing else; IPs and ports are "
       "carried only as analyst labels and are dropped by the ColumnTransformer before the model "
       "sees them — the design-level enforcement of the privacy guarantee in US-12. Table 4 is the "
       "feature schema; Figure 9 shows how each attack category expresses those features, which is "
       "why the trees can separate them.")
    _table(doc,
        ["Group", "Features"],
        [
            ["Volumetric", "flow_duration, total_fwd/bwd_packets, total_length_fwd/bwd_packets, flow_bytes_per_s, flow_packets_per_s"],
            ["Distributional", "fwd/bwd_packet_length_max/mean, flow_iat_mean, flow_iat_std, fwd/bwd_iat_total"],
            ["TCP flags", "fin/syn/rst/psh/ack_flag_count"],
            ["Protocol", "protocol (one-hot)"],
            ["Labels (dropped)", "src_ip, dst_ip, src_port, dst_port"],
        ],
        widths=[3.2, 11.8])
    _caption(doc, "Table 4 — The canonical feature schema (features/schema.py).")
    _figure(doc, "dataset.png",
            "Dataset class shares and per-category feature fingerprint — the signal the model keys "
            "on per attack class.")
    _p(doc,
       "Persistence is a single SQLite table written after each verdict. Table 5 is its shape; it "
       "is what backs the audit and investigation stories (US-03, US-04) and the post-incident "
       "replay runbook.")
    _table(doc,
        ["Column", "Type", "Purpose"],
        [
            ["id", "INTEGER PK", "Row identity."],
            ["ts", "TEXT (ISO)", "Classification timestamp (Amsterdam tz)."],
            ["verdict", "TEXT", "BENIGN or attack category."],
            ["score", "REAL", "Attack probability."],
            ["severity", "TEXT", "info / low / medium / high / critical."],
            ["src_ip / dst_ip", "TEXT", "Analyst labels (pseudonymisable)."],
            ["model", "TEXT", "Deciding model name (traceability)."],
            ["latency_ms", "REAL", "Per-request inference latency."],
        ],
        widths=[3.4, 2.6, 9.0])
    _caption(doc, "Table 5 — The predictions audit table (api/store.py).")

    _page_break(doc)

    # 7. API
    _h1(doc, "7. API specification")
    _p(doc,
       "The HTTP surface is intentionally small — six endpoints across four tags — and is fully "
       "OpenAPI-documented at /docs. Table 6 is the contract.")
    _table(doc,
        ["Method & path", "Tag", "Purpose", "Story"],
        [
            ["GET /health", "meta", "Model-load state + version.", "US-11"],
            ["GET /metrics", "meta", "Totals, severity, attacks-by-class, latency.", "US-11"],
            ["POST /predict", "inference", "Score a batch of flows; optional threshold.", "US-01, US-05"],
            ["POST /predict/csv", "inference", "Score an uploaded CICIDS-format CSV.", "US-01"],
            ["GET /alerts", "audit", "Severity-filtered alert feed.", "US-02, US-04"],
            ["GET /predictions", "audit", "Full audit log for triage.", "US-03"],
            ["POST /admin/reload", "admin", "Hot-reload a new model bundle.", "US-07"],
        ],
        widths=[4.6, 2.0, 6.0, 2.4])
    _caption(doc, "Table 6 — API surface mapped to user stories.")
    _p(doc,
       "A representative request and response for the core scoring path (US-01) — a batch of one "
       "flow, scored at the default operating point:")
    _code(doc,
       'POST /predict\n'
       '{ "flows": [ { "protocol": "TCP", "flow_duration": 120000,\n'
       '               "syn_flag_count": 1, "flow_packets_per_s": 50, ... } ],\n'
       '  "threshold": 0.5 }\n\n'
       '200 OK\n'
       '{ "results": [ { "verdict": "Infiltration", "score": 0.812,\n'
       '                 "severity": "high", "model": "gradient_boosting",\n'
       '                 "probabilities": { "Infiltration": 0.81, "BENIGN": 0.14, ... } } ] }')

    # 8. Runtime behaviour
    _h1(doc, "8. Runtime behaviour")
    _p(doc,
       "The most instructive runtime path is the attack/defend loop, because it exercises every "
       "component end to end. Figure 10 is the sequence: a click on the frontend drives the "
       "attacker container, real packets cross the bridge to the defender, the defender extracts a "
       "flow and posts it to /predict, the model returns a verdict, and on a confident attack the "
       "defender installs an iptables block — while the verdict simultaneously surfaces as a "
       "severity-coloured row in the analyst's Alerts tab. This single diagram realises US-09 and "
       "US-10 and demonstrates US-01 through US-03 in passing.")
    _figure(doc, "sequence.png",
            "Sequence — one attack through the detect-and-block pipeline: real tool → captured flow "
            "→ /predict verdict → automatic iptables block → analyst alert.")
    _p(doc,
       "The ordinary analyst path is a strict subset of the above: a client posts a batch to "
       "/predict, the FeaturePipeline transforms it, the model emits per-class probabilities, the "
       "service derives a severity band, writes the audit row asynchronously, and returns the "
       "verdict. Because the audit write is off the critical path, a slow disk degrades durability "
       "of the log, never the latency of the verdict — the design choice behind US-01's "
       "sub-50-millisecond acceptance criterion.")

    # 9. NFR & security
    _h1(doc, "9. Non-functional requirements and security")
    _p(doc,
       "Table 7 records the non-functional targets and where each is enforced. Security is treated "
       "as a design input, not an afterthought: the detector is itself threat-modelled with STRIDE "
       "in the Research Report, and the two controls most visible in this design are the "
       "no-payload feature schema (US-12) and the non-root, read-only-model container posture.")
    _table(doc,
        ["NFR", "Target", "Enforced by"],
        [
            ["Performance", "Sub-50 ms verdict per analyst batch", "Stateless /predict, batch-first, off-path audit"],
            ["Monitoring", "Every classification observable", "/metrics, /health, SQLite audit"],
            ["Security", "No PII in the inference path", "schema.py, ColumnTransformer(remainder=drop)"],
            ["Compliance", "GDPR/AVG: pseudonymisable, on-host, revocable", "IPs not features; delete alerts.db"],
            ["Reliability", "No-downtime model swap", "/admin/reload + ModelRegistry"],
            ["Usability", "Triage at a glance", "Severity colour-coding, audit beside queue"],
        ],
        widths=[2.6, 6.0, 6.4])
    _caption(doc, "Table 7 — Non-functional requirements and their enforcement points.")

    # 10. Deployment
    _h1(doc, "10. Deployment view")
    _p(doc,
       "Everything ships as Docker images orchestrated by Compose. The API and dashboard are split "
       "into two images so they can scale independently and keep small; both run as dedicated "
       "non-root users and mount the model directory read-only. The lab is a separate Compose file "
       "with two privileged-capability containers (NET_RAW for capture, NET_ADMIN for iptables) "
       "scoped to their own network namespace so the host firewall is never touched — the trust "
       "boundary of the offensive stories US-09 and US-10.")
    _bullets(doc, [
        "**ad-api** — FastAPI + model bundle, port 8000, non-root user detector.",
        "**ad-dashboard** — Streamlit, port 8501, non-root user dash.",
        "**React SPA** — built static assets served by nginx (prod) or Vite (dev), proxying /api.",
        "**ad-attacker / ad-defender** — the lab, brought up on demand from lab/docker-compose.yml.",
    ])

    # 11. Testing
    _h1(doc, "11. Testing strategy")
    _table(doc,
        ["Layer", "What is tested", "Where"],
        [
            ["Unit — features", "Pipeline robustness to NaN/inf, RobustScaler behaviour.", "tests/test_features.py"],
            ["Unit — API", "Happy path, attack flow, threshold, 422 on bad input.", "tests/ (API surface)"],
            ["Integration", "Train → serve → smoke-test on a small dataset.", "CI pipeline"],
            ["Model eval", "Per-class precision/recall/F1, confusion matrix.", "scripts/train, per_class_report.json"],
            ["Lab (manual)", "Run-on-lab → verdict → iptables block lands.", "lab/ + /blocks endpoint"],
        ],
        widths=[2.8, 7.4, 4.8])
    _caption(doc, "Table 8 — Test strategy across the layers.")

    # 12. Traceability
    _h1(doc, "12. Traceability matrix")
    _p(doc,
       "The matrix closes the loop from need to verification: every user story maps to the "
       "component that realises it, the endpoint that exposes it, and the test or evidence that "
       "confirms it. It is the single artefact a reviewer can use to check that nothing in the "
       "backlog is unbuilt and nothing built is unjustified.")
    _table(doc,
        ["Story", "Component", "Endpoint / surface", "Verified by"],
        [
            ["US-01", "FastAPI + FeaturePipeline", "POST /predict", "API tests; latency benchmark"],
            ["US-02", "Severity logic + dashboards", "/alerts", "API tests; dashboard"],
            ["US-03", "AlertStore", "GET /predictions", "API tests; audit table"],
            ["US-04", "AlertStore + dashboards", "/alerts rows", "Dashboard review"],
            ["US-05", "Threshold param + slider", "POST /predict?threshold", "API tests; manual scoring"],
            ["US-06", "Manual scoring view", "POST /predict", "Frontend; class-prob chart"],
            ["US-07", "ModelRegistry", "POST /admin/reload", "Reload smoke test"],
            ["US-08", "Training + evaluation", "scripts/train", "per_class_report.json"],
            ["US-09", "Lab attacker", "POST :8001/attack", "Manual lab run"],
            ["US-10", "Lab defender", "iptables + /blocks", "/blocks shows DROP"],
            ["US-11", "Metrics/health", "GET /health, /metrics", "API tests"],
            ["US-12", "FeaturePipeline schema", "(all)", "schema.py; no payload stored"],
        ],
        widths=[1.6, 4.6, 4.6, 4.2])
    _caption(doc, "Table 9 — Traceability: story → component → endpoint → verification.")

    _page_break(doc)
    _h1(doc, "References")
    for i, ref in enumerate([
        "Research Report — Network Traffic Anomaly Detector (companion document).",
        "C4 model for software architecture. https://c4model.com/",
        "Sharafaldin, I., Lashkari, A. H., & Ghorbani, A. A. (2018). Toward Generating a New Intrusion Detection Dataset. ICISSP.",
        "FastAPI documentation. https://fastapi.tiangolo.com/",
        "scikit-learn: Novelty and Outlier Detection. https://scikit-learn.org/stable/modules/outlier_detection.html",
        "OWASP API Security Top 10 (2023). https://owasp.org/API-Security/",
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.add_run(f"[{i}] ").bold = True
        p.add_run(ref)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT)
        print(f"[+] Wrote {OUT} with {_FIG} figures")
    except PermissionError:
        alt = OUT.with_name("Technical_Design_Document_new.docx")
        doc.save(alt)
        print(f"[!] {OUT.name} locked; wrote {alt} with {_FIG} figures")


def _code(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x22, 0x33, 0x44)


if __name__ == "__main__":
    build()

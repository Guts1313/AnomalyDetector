"""In-place surgery on the *styled* Research_Report_expanded.docx.

Non-destructive: keeps the user's manual styling. Two operations only:
  1. Replace the two C4 figures (Figure 2 context, Figure 3 container) with the
     new dark, grid-styled renders — overwriting the embedded image bytes and
     fixing each picture's display height for the new aspect ratio.
  2. Append a "Dark theme" gallery (Figures 20-23) at the end so that NO existing
     figure number shifts.

    python -m scripts.swap_doc_images
"""
from __future__ import annotations

import io
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "docs" / "Research_Report_expanded.docx"
SHOTS = ROOT / "docs" / "screenshots"

WIDTH_EMU = Cm(15.0).emu  # the report's standard figure width


def _replace_image(doc: Document, shape_index: int, png: Path) -> None:
    sh = doc.inline_shapes[shape_index]
    blip = sh._inline.graphic.graphicData.pic.blipFill.blip
    rid = blip.get(qn("r:embed"))
    part = doc.part.related_parts[rid]
    data = png.read_bytes()
    part._blob = data
    w, h = Image.open(io.BytesIO(data)).size
    sh.width = Emu(WIDTH_EMU)
    sh.height = Emu(round(WIDTH_EMU * h / w))
    print(f"  [+] shape {shape_index} <- {png.name} ({w}x{h}) @ {WIDTH_EMU/360000:.1f}cm")


def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x65, 0x70)
    p.paragraph_format.space_after = Pt(12)


def _figure(doc: Document, png: Path, caption: str, width_cm: float = 15.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png), width=Cm(width_cm))
    _caption(doc, caption)


def main() -> None:
    doc = Document(str(DOC))

    print("[*] Replacing C4 figures in place:")
    _replace_image(doc, 2, SHOTS / "21_c4_context.png")     # Figure 2 — context
    _replace_image(doc, 3, SHOTS / "22_c4_container.png")   # Figure 3 — container

    print("[*] Appending dark-theme gallery (Figures 20-23):")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Appendix E — The dark theme", level=1)
    doc.add_paragraph(
        "The React interface ships in two themes and remembers the analyst's choice. The figures "
        "in the body of this report use the light (daylight) theme; the gallery below shows the "
        "same screens in the dark theme a SOC tends to run on a wall display, where a bright screen "
        "is fatiguing over a long shift. The theme is a single toggle in the navigation bar and is "
        "purely cosmetic — the severity colour-coding, the charts and the audit semantics are "
        "identical across both, so nothing in the analysis changes with it."
    )
    _figure(doc, SHOTS / "30_fe_overview_dark.png",
            "Figure 20 — React overview in the dark theme: the same KPIs and live severity / "
            "attacks-by-class charts against the deep-blue glass surface.")
    _figure(doc, SHOTS / "31_fe_alerts_dark.png",
            "Figure 21 — Dark-theme alerts table: severity badges and glowing attack-score columns "
            "over the audit log.")
    _figure(doc, SHOTS / "32_fe_manual_dark.png",
            "Figure 22 — Dark-theme manual scoring: the same hand-built flow returns an infiltration "
            "verdict and a full class-probability chart.", width_cm=14.5)
    _figure(doc, SHOTS / "33_fe_examples_dark.png",
            "Figure 23 — Dark-theme Request-examples tab: the attack/defend control plane, the "
            "active category framed by the animated accent border.", width_cm=14.5)

    doc.save(str(DOC))
    print(f"[+] Saved {DOC}")


if __name__ == "__main__":
    main()
